# -*- coding: utf-8 -*-
"""GB/T 9704-2012《党政机关公文格式》套版核心(python-docx)。

把结构化公文 JSON 渲染成合规 .docx:字体/字号/页边距/行距/版心都按国标硬编码。
内容(话术/序号)由模型经 skill 产出进 JSON;此处只管"套版",一个字不改写。

字号(中文号→pt,GB/T 9704 常用):二号=22,三号=16,小三=15,四号=14。
字体:标题=方正小标宋简体;一级"一、"=黑体;二级"（一）"=楷体_GB2312;
      三级"1."=仿宋_GB2312 加粗;正文/发文字号/主送/落款=仿宋_GB2312。
中文字体名会写进 docx 的 w:eastAsia,在装了对应字体的 Word/WPS 里正确呈现;
开发机不装这些字体不影响 .docx 文件本身的正确性(打开方按字体名渲染)。
"""
import re

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 字号(pt)──────────────────────────────────────────────────────────────
SZ_ERHAO = 22      # 二号:文件标题
SZ_SANHAO = 16     # 三号:正文及绝大多数部位
SZ_REDHEAD = 28    # 简化红头(发文机关标志)用大字,仅观感

# ── 字体名 ────────────────────────────────────────────────────────────────
F_BIAOTI = "方正小标宋简体"   # 标题
F_HEITI = "黑体"             # 一级标题"一、"
F_KAITI = "楷体_GB2312"      # 二级标题"（一）"
F_FANGSONG = "仿宋_GB2312"   # 正文/三级/发文字号/主送/落款
F_LATIN = "Times New Roman"  # 数字与英文

# 正文级别 → (字体, 是否加粗)。"正文"另加首行缩进 2 字。
LEVEL_FONT = {
    "一级": (F_HEITI, False),
    "二级": (F_KAITI, False),
    "三级": (F_FANGSONG, True),
    "四级": (F_FANGSONG, False),
    "正文": (F_FANGSONG, False),
}


def _set_cjk_font(run, cjk_name, size_pt, bold=False, color=None):
    """python-docx 设中文字体必须同时写 w:eastAsia,否则中文回退默认宋体。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = F_LATIN  # 拉丁/数字
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cjk_name)
    rfonts.set(qn("w:ascii"), F_LATIN)
    rfonts.set(qn("w:hAnsi"), F_LATIN)


def _para(doc, text, cjk, size, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent_chars=0, color=None, space_after=0, line_pt=28):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)          # GB/T 9704:正文固定行距约 28–29 磅
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if first_indent_chars:
        pf.first_line_indent = Pt(first_indent_chars * size)  # n 个字宽
    run = p.add_run(text)
    _set_cjk_font(run, cjk, size, bold=bold, color=color)
    return p


def _red_separator(doc):
    """红色分隔线(版头与正文之间)。用段落下边框实现。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")        # 线粗
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF0000") # 红
    pbdr.append(bottom)
    pPr.append(pbdr)


# 序号前缀 → 级别(确定性纠偏:模型偶发把『二、』标成二级致字体打架,按前缀强制对齐)
_RE_CHAPTER = re.compile(r"^第[一二三四五六七八九十百零〇]+[章节]")  # 第一章/第二节(法规体例)
_RE_L1 = re.compile(r"^[一二三四五六七八九十]+、")          # 一、总体要求
_RE_L2 = re.compile(r"^[（(][一二三四五六七八九十]+[）)]")    # （一）…
_RE_L3 = re.compile(r"^\d+[.．、]")                          # 1. / 1、


def _infer_level(text, declared):
    """『级别』优先由序号前缀确定(第N章=章 / 一、=一级 / （一）=二级 / 1.=三级),
       前缀不明确时才尊重模型自报的 declared,缺省『正文』。
       『第N条』不特判:它本就是仿宋正文(首行缩进2字),走缺省即可。"""
    t = text.lstrip()
    if _RE_CHAPTER.match(t):
        return "章"
    if _RE_L1.match(t):
        return "一级"
    if _RE_L2.match(t):
        return "二级"
    if _RE_L3.match(t):
        return "三级"
    return declared or "正文"


def _render_body(doc, blocks):
    """渲染正文 block 数组:每项 {级别, 文字}。"""
    for b in blocks or []:
        text = (b.get("文字") or "").strip()
        if not text:
            continue
        level = _infer_level(text, (b.get("级别") or "").strip())
        if level == "章":  # 法规体例「第N章」:黑体三号·居中(区别于条文正文)
            _para(doc, text, F_HEITI, SZ_SANHAO,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
            continue
        font, bold = LEVEL_FONT.get(level, (F_FANGSONG, False))
        indent = 2 if level == "正文" else 0
        _para(doc, text, font, SZ_SANHAO, bold=bold,
              first_indent_chars=indent, align=WD_ALIGN_PARAGRAPH.LEFT)


def render_gongwen(data, out_path):
    """把公文 JSON 渲染成 GB/T 9704 .docx,落 out_path。返回写出的段落数(粗略)。"""
    doc = Document()

    # 版心:A4 + 国标页边距(上37 下35 左28 右26 mm)
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin, sec.bottom_margin = Mm(37), Mm(35)
    sec.left_margin, sec.right_margin = Mm(28), Mm(26)

    fwjg = (data.get("发文机关") or "").strip()
    fwzh = (data.get("发文字号") or "").strip()

    # ── 简化版头:红头机关名(大字红色居中)+ 发文字号 + 红色分隔线 ──
    if fwjg:
        _para(doc, fwjg, F_BIAOTI, SZ_REDHEAD, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=RGBColor(0xFF, 0x00, 0x00), space_after=6, line_pt=34)
    if fwzh:
        _para(doc, fwzh, F_FANGSONG, SZ_SANHAO, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _red_separator(doc)

    # ── 标题(方正小标宋二号·居中)──
    title = (data.get("标题") or "").strip()
    if title:
        _para(doc, title, F_BIAOTI, SZ_ERHAO, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=10, line_pt=34)

    # ── 主送机关(仿宋三号·顶格·末尾冒号）──
    zhusong = (data.get("主送机关") or "").strip()
    if zhusong:
        if not zhusong.endswith("：") and not zhusong.endswith(":"):
            zhusong += "："
        _para(doc, zhusong, F_FANGSONG, SZ_SANHAO, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ── 正文 ──
    _render_body(doc, data.get("正文"))

    # ── 附件(印发型通知的主件:换页 + 居中标题 + 各自正文)──
    for att in data.get("附件") or []:
        doc.add_page_break()
        at = (att.get("标题") or "").strip()
        if at:
            _para(doc, at, F_BIAOTI, SZ_ERHAO, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=10, line_pt=34)
        _render_body(doc, att.get("正文"))

    # ── 落款机关 + 成文日期(右对齐·仿宋三号)──
    luokuan = (data.get("落款机关") or fwjg or "").strip()
    riqi = (data.get("成文日期") or "").strip()
    if luokuan or riqi:
        doc.add_paragraph()  # 与正文留一空行
        if luokuan:
            _para(doc, luokuan, F_FANGSONG, SZ_SANHAO, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if riqi:
            _para(doc, riqi, F_FANGSONG, SZ_SANHAO, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── 版记:公开方式 + 印发说明 ──
    gongkai = (data.get("公开方式") or "主动公开").strip()
    doc.add_paragraph()
    _para(doc, "公开方式：" + gongkai, F_FANGSONG, SZ_SANHAO, align=WD_ALIGN_PARAGRAPH.LEFT)
    yinfa = (data.get("印发说明") or "").strip()
    if yinfa:
        _para(doc, yinfa, F_FANGSONG, SZ_SANHAO, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(out_path)
    return len(doc.paragraphs)
